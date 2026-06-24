import os
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- XML Helper Functions for Advanced formatting ---

def set_cell_background(cell, color_hex):
    """Set background color of a cell (e.g. 'F0F0F0' for light grey)."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Check if shading already exists
    shd = tcPr.find(qn('w:shading'))
    if shd is not None:
        tcPr.remove(shd)
    shading = OxmlElement('w:shading')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    tcPr.append(shading)

def add_page_border(section, color="000000", sz="4", space="20"):
    """Add a professional thin page border around the section at the correct XML schema position."""
    secPr = section._sectPr
    
    # Check if pgBorders already exists
    pgBorders = secPr.find(qn('w:pgBorders'))
    if pgBorders is not None:
        secPr.remove(pgBorders)
        
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(sz))
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), color)
        pgBorders.append(border)
    
    # Correct OpenXML ordering: pgBorders goes before lnNumType, pgNumType, cols, docGrid
    after_tags = [qn('w:lnNumType'), qn('w:pgNumType'), qn('w:cols'), qn('w:docGrid'), qn('w:printerSettings')]
    inserted = False
    for child in secPr:
        if child.tag in after_tags:
            child.addprevious(pgBorders)
            inserted = True
            break
            
    if not inserted:
        secPr.append(pgBorders)

def add_page_number(paragraph):
    """Insert a dynamic PAGE number field in separate runs inside a paragraph."""
    paragraph.text = ""
    
    r_begin = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fldChar1)
    
    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    r_instr.append(instrText)
    
    r_sep = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fldChar2)
    
    # Placeholder number
    r_num = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "1"
    r_num.append(t)
    
    r_end = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_end.append(fldChar3)
    
    p = paragraph._p
    p.append(r_begin)
    p.append(r_instr)
    p.append(r_sep)
    p.append(r_num)
    p.append(r_end)

def add_toc(paragraph):
    """Insert an automatic Table of Contents field block at the correct run-level sequence."""
    paragraph.text = ""
    run = paragraph.add_run("CONTENT\n\n")
    run.bold = True
    run.font.size = Pt(16)
    
    # Label headers
    p_labels = paragraph.insert_paragraph_before()
    p_labels.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_labels.paragraph_format.space_after = Pt(6)
    
    run_lbl1 = p_labels.add_run("Content")
    run_lbl1.bold = True
    run_lbl1.font.size = Pt(12)
    
    # Use tabs for right alignment
    p_labels.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=2) # 2 = right
    run_lbl2 = p_labels.add_run("\tPage No.")
    run_lbl2.bold = True
    run_lbl2.font.size = Pt(12)
    
    # Wrap TOC in separate runs
    r_begin = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fldChar1)
    
    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    r_instr.append(instrText)
    
    r_sep = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fldChar2)
    
    r_text = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "[Right-click here and select 'Update Field' to generate Table of Contents]"
    r_text.append(t)
    
    r_end = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_end.append(fldChar3)
    
    p = paragraph._p
    p.append(r_begin)
    p.append(r_instr)
    p.append(r_sep)
    p.append(r_text)
    p.append(r_end)

def set_page_numbering(section, format_name="decimal", start_val=None):
    """Configure page number display properties at the correct XML schema position."""
    secPr = section._sectPr
    
    pgNumType = secPr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        secPr.remove(pgNumType)
        
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:fmt'), format_name)
    if start_val is not None:
        pgNumType.set(qn('w:start'), str(start_val))
        
    after_tags = [qn('w:cols'), qn('w:docGrid'), qn('w:printerSettings')]
    inserted = False
    for child in secPr:
        if child.tag in after_tags:
            child.addprevious(pgNumType)
            inserted = True
            break
            
    if not inserted:
        secPr.append(pgNumType)

# --- Document Generation Setup ---

doc = Document()

# Set standard 1-inch margins on all sides
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

cover_section = doc.sections[0]
add_page_border(cover_section, sz="6", space="20")

# Setup default fonts and styles
styles = doc.styles

# Configure Normal style
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Times New Roman'
normal_font.size = Pt(12)
normal_font.color.rgb = RGBColor(0, 0, 0)
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal_style.paragraph_format.space_after = Pt(6)
normal_style.paragraph_format.space_before = Pt(0)

# Configure headings
def setup_heading_style(name, size, space_before=12, space_after=6):
    style = styles[name]
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(size)
    font.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

setup_heading_style('Heading 1', 16, space_before=18, space_after=8)
setup_heading_style('Heading 2', 14, space_before=14, space_after=6)
setup_heading_style('Heading 3', 12, space_before=12, space_after=6)

# --- Download the Logo ---
logo_path = "gm_university_logo.png"
logo_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/4/4c/GM_University%2C_Davanagere.png/1280px-GM_University%2C_Davanagere.png"

has_logo = False
try:
    print(f"Downloading logo from {logo_url}...")
    req = urllib.request.Request(
        logo_url, 
        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    )
    with urllib.request.urlopen(req) as response, open(logo_path, 'wb') as out_file:
        out_file.write(response.read())
    has_logo = True
    print("Logo downloaded successfully.")
except Exception as e:
    print(f"Failed to download logo: {e}. Will use a styled text placeholder.")

# ==========================================
# SECTION 1: COVER PAGE
# ==========================================

# 1. Srishyla Educational Trust ®
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = p1.add_run("Srishyla Educational Trust ®")
run1.font.size = Pt(12)
run1.font.name = 'Times New Roman'
p1.paragraph_format.space_before = Pt(0)
p1.paragraph_format.space_after = Pt(2)

# 2. GM UNIVERSITY
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("GM UNIVERSITY")
run2.font.size = Pt(28)
run2.bold = True
run2.font.name = 'Times New Roman'
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(2)

# 3. Established details
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("(Established under the Karnataka State Act No. 19 of 2023) Post Box No. 4, Davanagere-577006")
run3.font.size = Pt(9.5)
run3.font.name = 'Times New Roman'
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(18)

# 4. Logo Image
if has_logo:
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(1.8))
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(18)
else:
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("\n[ GM UNIVERSITY LOGO ]\n")
    run_logo.font.size = Pt(14)
    run_logo.bold = True
    run_logo.font.name = 'Times New Roman'
    p_logo.paragraph_format.space_before = Pt(12)
    p_logo.paragraph_format.space_after = Pt(30)

# 5. Faculty & Department details
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("Faculty of Engineering and Technology")
run4.font.size = Pt(15)
run4.bold = True
run4.font.name = 'Times New Roman'
p4.paragraph_format.space_after = Pt(4)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING")
run5.font.size = Pt(14)
run5.bold = True
run5.font.name = 'Times New Roman'
p5.paragraph_format.space_after = Pt(18)

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run("Academic Year 2025 - 26")
run6.font.size = Pt(13)
run6.italic = True
run6.bold = True
run6.font.name = 'Times New Roman'
p6.paragraph_format.space_after = Pt(14)

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run7 = p7.add_run("Project Based Learning")
run7.font.size = Pt(16)
run7.bold = True
run7.font.name = 'Times New Roman'
p7.paragraph_format.space_after = Pt(4)

p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
run8 = p8.add_run("AI BASED RESUME SCREENING SYSTEM")
run8.font.size = Pt(18)
run8.bold = True
run8.font.name = 'Times New Roman'
p8.paragraph_format.space_after = Pt(4)

p9 = doc.add_paragraph()
p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
run9 = p9.add_run("Semester: 4th")
run9.font.size = Pt(14)
run9.bold = True
run9.font.name = 'Times New Roman'
p9.paragraph_format.space_after = Pt(14)

# 6. Submitted By: Table
p_sb = doc.add_paragraph()
p_sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sb = p_sb.add_run("Submitted By:")
run_sb.font.size = Pt(12)
run_sb.bold = True
p_sb.paragraph_format.space_after = Pt(6)

# Use standard 'Table Grid' to ensure borders are loaded cleanly in MS Word
table_sb = doc.add_table(rows=5, cols=2)
table_sb.style = 'Table Grid'
table_sb.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set Table Columns width
widths = [Inches(2.5), Inches(2.0)]
for row in table_sb.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w

# Format headers
hdr_cells = table_sb.rows[0].cells
hdr_cells[0].paragraphs[0].text = ""
hdr_run1 = hdr_cells[0].paragraphs[0].add_run("Name")
hdr_run1.bold = True
hdr_run1.font.size = Pt(11)
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[0].paragraphs[0].paragraph_format.space_before = Pt(4)
hdr_cells[0].paragraphs[0].paragraph_format.space_after = Pt(4)

hdr_cells[1].paragraphs[0].text = ""
hdr_run2 = hdr_cells[1].paragraphs[0].add_run("USN")
hdr_run2.bold = True
hdr_run2.font.size = Pt(11)
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[1].paragraphs[0].paragraph_format.space_before = Pt(4)
hdr_cells[1].paragraphs[0].paragraph_format.space_after = Pt(4)

# Set shading for headers
set_cell_background(hdr_cells[0], "E0E0E0")
set_cell_background(hdr_cells[1], "E0E0E0")

# Adjust spacing for empty rows so they are visually padded
for row_idx in range(1, 5):
    for cell in table_sb.rows[row_idx].cells:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

# Space before Submitted To
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(12)
p_space.paragraph_format.space_after = Pt(2)

# 7. Submitted To: Section
p_st = doc.add_paragraph()
p_st.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_st = p_st.add_run("Submitted To:")
run_st.font.size = Pt(12)
run_st.bold = True
p_st.paragraph_format.space_after = Pt(14)

# Three-column borderless table for Submitted To layout
table_st = doc.add_table(rows=1, cols=3)
table_st.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Inches(2.2), Inches(2.2), Inches(2.2)]
for cell_idx, cell in enumerate(table_st.rows[0].cells):
    cell.width = col_widths[cell_idx]

# Column 1: Guide
cell_guide = table_st.rows[0].cells[0]
p_g = cell_guide.paragraphs[0]
p_g.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_g.paragraph_format.line_spacing = 1.15
p_g.paragraph_format.space_after = Pt(0)
p_g.add_run("_____________________\n").bold = True
p_g.add_run("Guide\n").bold = True
p_g.add_run("(Name)\n").bold = True
p_g.add_run("(Designation)\n").bold = True
p_g.add_run("GM University").bold = True

# Column 2: PBL Coordinator
cell_pbl = table_st.rows[0].cells[1]
p_p = cell_pbl.paragraphs[0]
p_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_p.paragraph_format.line_spacing = 1.15
p_p.paragraph_format.space_after = Pt(0)
p_p.add_run("_____________________\n").bold = True
p_p.add_run("PBL Coordinator\n").bold = True
p_p.add_run("Mrs. Sindhu R R\n").bold = True
p_p.add_run("Tutor\n").bold = True
p_p.add_run("GM University").bold = True

# Column 3: HOD
cell_hod = table_st.rows[0].cells[2]
p_h = cell_hod.paragraphs[0]
p_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_h.paragraph_format.line_spacing = 1.15
p_h.paragraph_format.space_after = Pt(0)
p_h.add_run("_____________________\n").bold = True
p_h.add_run("HOD\n").bold = True
p_h.add_run("Dr. Shivanagowda G M\n").bold = True
p_h.add_run("Professor & Head\n").bold = True
p_h.add_run("GM University").bold = True


# ==========================================
# SECTION 2: ABSTRACT & TOC (ROMAN NUMERALS)
# ==========================================

section_tocs = doc.add_section()
section_tocs.top_margin = Inches(1.0)
section_tocs.bottom_margin = Inches(1.0)
section_tocs.left_margin = Inches(1.0)
section_tocs.right_margin = Inches(1.0)

# Configure page numbering to Roman numerals
set_page_numbering(section_tocs, format_name="romanLower", start_val=1)

section_tocs.header.is_linked_to_previous = False
section_tocs.footer.is_linked_to_previous = False

# Add page number to footer
footer_p = section_tocs.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(footer_p)

# --- Table of Contents Page ---
p_toc_title = doc.add_paragraph()
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_toc_title.paragraph_format.space_before = Pt(24)
p_toc_title.paragraph_format.space_after = Pt(24)
add_toc(p_toc_title)

# --- Abstract Page ---
doc.add_page_break()

p_abs = doc.add_heading("Abstract", level=1)
p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_abs.paragraph_format.space_before = Pt(24)
p_abs.paragraph_format.space_after = Pt(18)

p_abs_text = doc.add_paragraph(
    "In the modern recruitment industry, screening resumes is a major bottleneck due to the high volume of "
    "applications received for each job opening. Traditional recruitment methods are time-consuming, prone to human "
    "cognitive bias, and highly inefficient. To address these challenges, this project presents the design and "
    "implementation of the AI Resume Screening System (ARSS), an automated platform that evaluates candidate resumes "
    "against job requirements. The core AI pipeline consists of a five-step process: raw text extraction, structured "
    "information extraction, semantic similarity matching, hybrid scoring, and classification. Using PyMuPDF and python-docx, "
    "the system extracts raw text from PDF and DOCX files. Natural Language Processing (NLP) models, specifically spaCy's "
    "Named Entity Recognition (NER), are employed to identify the candidate's name. Email, phone number, and years of experience "
    "are extracted using regular expressions, while skills and education are parsed using keyword-matching logic. "
    "To measure document alignment, the system uses Term Frequency-Inverse Document Frequency (TF-IDF) vectorization "
    "and cosine similarity to compare the parsed resume text with the job description. The final screening score is computed "
    "using a hybrid formula that weights the semantic similarity score at 70% and the normalized years of experience at 30%. "
    "Education serves as a hard gate: candidates with mismatched qualifications are immediately rejected. The final score "
    "determines whether the candidate is classified as QUALIFIED, SHORTLIST, or REJECT. "
    "Furthermore, a persistent HTTP ML server architecture is implemented to optimize the processing pipeline, reducing "
    "latency from 15 seconds to under 500 milliseconds per resume. The resulting system is hosted on a React frontend and a "
    "Node.js/Express backend with MongoDB, offering recruiter dashboards and candidate interfaces. The final system "
    "provides an objective, highly responsive, and robust solution for modern talent acquisition."
)


# ==========================================
# SECTION 3: MAIN CHAPTERS (ARABIC NUMERALS)
# ==========================================

section_main = doc.add_section()
section_main.top_margin = Inches(1.0)
section_main.bottom_margin = Inches(1.0)
section_main.left_margin = Inches(1.0)
section_main.right_margin = Inches(1.0)

# Configure page numbering to standard decimals
set_page_numbering(section_main, format_name="decimal", start_val=1)

section_main.header.is_linked_to_previous = False
section_main.footer.is_linked_to_previous = False

footer_main = section_main.footer.paragraphs[0]
footer_main.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_page_number(footer_main)

# --- CHAPTER 1 ---
doc.add_heading("CHAPTER 1 . Introduction", level=1)

doc.add_paragraph(
    "The hiring process in modern organizations involves managing hundreds or thousands of resumes for a single position. "
    "Manual resume screening is one of the most tedious tasks in human resource management. Recruiters spend an average of "
    "six to ten seconds reviewing a single resume, which can lead to fatigue, oversight of qualified candidates, and "
    "inherent human bias. In recent years, Applicant Tracking Systems (ATS) have emerged as software solutions to filter and "
    "organize candidate profiles. However, traditional ATS platforms rely on rigid keyword searches, which often fail to "
    "understand semantic context, rejecting strong candidates who use synonyms and passing weaker candidates who keyword-stuff "
    "their resumes."
)

doc.add_paragraph(
    "The AI Resume Screening System (ARSS) developed in this project addresses these limitations. ARSS is a full-stack, "
    "intelligent application that automates the screening process. It integrates a React-based client interface with a "
    "Node.js backend, backed by MongoDB for persistence, and bridges to a high-performance Python-based Machine Learning (ML) "
    "pipeline. The Python pipeline extracts and structures resume content, computes semantic similarity against job requirements, "
    "and classifies candidates using objective scoring metrics. To make this pipeline production-ready, the ML models are "
    "hosted on a persistent, localized HTTP server, loading heavy NLP libraries once at boot to process individual resumes in "
    "real-time (200-500ms)."
)

# --- CHAPTER 2 ---
doc.add_heading("CHAPTER 2 . Background Literature", level=1)

doc.add_paragraph(
    "Automated resume parsing and matching has evolved through several distinct phases. Initial attempts relied on simple "
    "rule-based heuristics and regular expressions. While effective for highly structured files, these early systems were "
    "extremely fragile when faced with creative layouts, columns, or varied styling. The advent of Natural Language "
    "Processing (NLP) and Named Entity Recognition (NER) marked a major shift. NER models, such as those provided by the spaCy "
    "framework, use convolutional neural networks trained on large corpora to recognize semantic entities like names, "
    "organizations, and locations. By utilizing spaCy's 'en_core_web_sm' model, the proposed system extracts candidate "
    "names with higher accuracy, regardless of their position or layout on the page."
)

doc.add_paragraph(
    "For document matching, early systems relied on exact keyword overlap. The modern approach utilizes vector space modeling. "
    "Term Frequency-Inverse Document Frequency (TF-IDF) is a statistical measure used to evaluate how important a word is "
    "to a document in a collection. When applied to job descriptions and resumes, TF-IDF represents text as high-dimensional "
    "vectors. The angle between these vectors is measured using Cosine Similarity, yielding a score between 0.0 and 1.0. This "
    "combination allows the system to determine semantic alignment, ensuring that the vocabulary used in the resume corresponds "
    "naturally to the requirements of the job description without requiring exact, word-for-word matching."
)

# --- CHAPTER 3 ---
doc.add_heading("CHAPTER 3 . Problem Statement and Objectives", level=1)

doc.add_heading("3.1 Problem Statement", level=2)
doc.add_paragraph(
    "Manual resume screening is a significant bottleneck in recruitment, characterized by high operational latency, "
    "susceptibility to unconscious human bias, and administrative overhead. Existing legacy Applicant Tracking Systems "
    "rely on exact keyword-matching rules, which reject qualified candidates due to lexical variation (e.g., using 'deep learning' "
    "instead of 'neural networks'). Conversely, these systems fail to extract structured data (such as years of experience "
    "and education levels) accurately from unstructured text, requiring recruiters to manually verify each candidate's "
    "basic eligibility. There is a critical need for an intelligent, low-latency, and objective system that can parse, "
    "extract structured entities, calculate semantic matching, and classify candidates with high responsiveness."
)

doc.add_heading("3.2 Objectives", level=2)
doc.add_paragraph(
    "To address the defined problem, the project outlines the following primary objectives:\n"
    "1. To develop a robust file-parsing module capable of extracting clean, raw text from PDF and DOCX formats using PyMuPDF and python-docx.\n"
    "2. To implement a structured entity extraction pipeline utilizing spaCy Named Entity Recognition for candidate names, and custom regular expressions for phone numbers, emails, and years of experience.\n"
    "3. To implement a semantic document matching model using TF-IDF vectorization and Cosine Similarity to evaluate candidate text against job descriptions.\n"
    "4. To design a hybrid scoring and classification engine that incorporates years of experience and treats education level as a hard gate.\n"
    "5. To optimize pipeline execution speed using a persistent HTTP ML server architecture to load model weights once and process resumes in under 500ms.\n"
    "6. To construct a responsive React frontend and Node.js backend with MongoDB, providing candidate uploading portals and recruiter-facing administration dashboards."
)

# --- CHAPTER 4 ---
doc.add_heading("CHAPTER 4 .Solution Methodology", level=1)

doc.add_paragraph(
    "The proposed Solution Methodology follows a structured, modular design divided into two primary parts: the Web Application "
    "Architecture and the core 5-Step AI Resume Screening Pipeline."
)

doc.add_paragraph(
    "The 5-Step AI Pipeline is coordinated by the 'pipeline_runner.py' script and executes the following sequence:\n"
    "1. Text Parsing: The input file (.pdf or .docx) is read. PyMuPDF handles PDF layouts, extracting characters page-by-page, while python-docx extracts paragraph strings from Word documents.\n"
    "2. Entity Extraction: The raw text is processed. Candidate name is extracted using spaCy's PERSON NER entity. Contact details are parsed via regex. Skills are identified via lookup tables, and years of experience are extracted using context-aware regular expressions (e.g., matching 'X+ years of experience').\n"
    "3. Semantic Similarity: The job description from the configuration is vectorized using TF-IDF. The candidate's resume text is transformed into the same vector space. Cosine similarity calculates the alignment between the two vectors.\n"
    "4. Hybrid Scoring: The final score is computed as: Score = (0.7 * Similarity) + (0.3 * Normalized Experience). Experience is capped at 5 years and normalized by dividing by 5. This weights keywords heavily while giving an objective advantage to experienced applicants.\n"
    "5. Hard-Gate Classification: The candidate's extracted education level (e.g., 'B.TECH') is compared with the required education. If they do not match, the candidate is classified as 'REJECT' immediately. Otherwise, the candidate is classified based on the hybrid score: 'QUALIFIED' (>= 0.75), 'SHORTLIST' (0.40 - 0.74), or 'REJECT' (< 0.40)."
)

# --- CHAPTER 5 ---
doc.add_heading("CHAPTER 5. Implementation", level=1)

doc.add_paragraph(
    "The system is implemented as a full-stack JavaScript and Python application. The repository is organized into three major divisions: "
    "the Node.js/Express server, the React/Vite client, and the Python AI pipeline modules."
)

doc.add_heading("5.1 Backend and Database Setup", level=2)
doc.add_paragraph(
    "The backend is constructed using Node.js and Express.js, exposing RESTful API endpoints for candidate uploads, recruiter authentication, "
    "dashboard statistics, and requirement configurations. Recruiter accounts are protected using JSON Web Tokens (JWT) stored in secure cookies, "
    "and file uploads are handled via Multer middleware. MongoDB acts as the primary data store, using Mongoose schemas for data modeling. "
    "The 'Candidate' schema stores contact info, extracted skills, experience, education, computed similarity, overall score, final classification, "
    "and status (Applied, Shortlisted, Interviewing, Rejected, Hired). The 'Config' schema stores the active job requirements, including the job "
    "description, minimum experience, and required education level."
)

doc.add_heading("5.2 Python Pipeline Integration", level=2)
doc.add_paragraph(
    "To connect Node.js with the Python ML code, a persistent child-process bridge is established. Upon backend boot, Node.js starts "
    "the 'pipeline_runner.py' script in server mode, bound to port 5001. The Python script loads the spaCy model, scikit-learn classes, "
    "and PyMuPDF libraries into memory once. When a candidate uploads a resume, Node.js performs a local HTTP POST request to localhost:5001 "
    "containing the file path and active job configuration. The Python server processes the resume in-memory and returns a structured JSON payload, "
    "which is then saved to MongoDB."
)

doc.add_heading("5.3 Frontend Interface", level=2)
doc.add_paragraph(
    "The user interface is built as a single-page application using React, Vite, and Redux Toolkit for global state management. "
    "It features a clean Landing Page, a Candidate Upload Page with real-time progress bars, and a Recruiter Dashboard. The dashboard "
    "is restricted to authenticated admins, displaying analytics cards for active applications and a paginated table of candidates. "
    "Recruiters can modify the job requirements (updating the job description, skills, and education level in the database), which immediately "
    "updates the screening criteria for subsequent resume uploads."
)

# --- CHAPTER 6 ---
doc.add_heading("CHAPTER 6. Presentation of Results", level=1)

doc.add_paragraph(
    "The system generates structured screening outputs that are rendered visually on both the candidate and recruiter views. "
    "Upon successful processing of a resume, the application outputs a detailed feedback object."
)

doc.add_paragraph(
    "For the recruiter, the candidate's profile is populated on the admin dashboard. The entry details the candidate's parsed name, "
    "email, phone number, calculated years of experience, and highest education. It displays the semantic similarity score, the "
    "hybrid score, and the final classification status (e.g., 'QUALIFIED' in a green tag, 'SHORTLIST' in yellow, and 'REJECT' in red). "
    "Recruiters can download the original resume directly from the dashboard."
)

doc.add_paragraph(
    "For the candidate, the portal displays a matching percentage based on the hybrid score. To ensure transparency, the system "
    "generates actionable resume improvement suggestions. By comparing the extracted skills against the required skills list, the pipeline "
    "identifies missing skills (e.g., 'Consider adding these missing skills: tensorflow, django'). It also flags structural weaknesses, "
    "such as a missing email address or low experience levels, and displays a categorized list of strengths (e.g., 'Proficient in Python') "
    "and weaknesses (e.g., 'Missing skill: SQL')."
)

# --- CHAPTER 7 ---
doc.add_heading("CHAPTER 7. Performance Analysis", level=1)

doc.add_paragraph(
    "Performance analysis was conducted to measure two primary aspects: the computational efficiency of the persistent HTTP pipeline "
    "server versus the traditional command-line interface (CLI) execution, and the accuracy of the parsing modules."
)

doc.add_heading("7.1 Computational Latency Analysis", level=2)
doc.add_paragraph(
    "Executing machine learning models in Python often incurs high initialization costs. When 'pipeline_runner.py' is executed in CLI mode, "
    "the interpreter must boot, import packages, load spaCy's 'en_core_web_sm' model (~12MB) and scikit-learn libraries, and compile C extensions. "
    "This cold-start process takes between 12 to 15 seconds per resume. In contrast, the persistent HTTP server mode loads all models "
    "and packages into RAM exactly once at backend startup. Subsequent requests bypass the cold start entirely. The average latency in server "
    "mode ranges from 200 to 500 milliseconds per resume, resulting in a 30x speedup and enabling high-throughput, concurrent screening."
)

doc.add_heading("7.2 Extractor Accuracy and Limitations", level=2)
doc.add_paragraph(
    "The extraction accuracy of the individual modules was benchmarked using sample resumes:\n"
    "1. Name Extraction: spaCy's NER successfully extracts standard English names located near the top of the resume. However, it displays limitations when processing highly stylized resumes, multi-column headers, or extremely rare names, which sometimes fail to trigger the 'PERSON' tag.\n"
    "2. Email and Phone: The regular expressions achieve 100% precision and recall for standard email and telephone formatting structures.\n"
    "3. Years of Experience: The experience regex accurately captures standard phrasing ('3 years of experience', '5+ yrs experience'). It occasionally misses experience if it is documented solely as date ranges (e.g., '2020 - 2023') without a summary count, identifying a future area of improvement."
)

# --- CHAPTER 8 ---
doc.add_heading("CHAPTER 8 . Applications", level=1)

doc.add_paragraph(
    "The AI Resume Screening System has direct applications in several areas of modern talent acquisition and human resources:\n"
    "1. Corporate Recruitment: HR departments can integrate ARSS into their careers portal. Large enterprises that receive thousands of applications weekly can instantly filter out unqualified profiles and highlight top talent, saving hundreds of recruiter hours.\n"
    "2. Recruitment Agencies: Third-party staffing firms can use the system to parse and index massive resume databases, allowing them to search and match candidate pools with new job descriptions in seconds.\n"
    "3. Educational Institutions: Placement cells in universities can use the candidate-facing portal to help students optimize their resumes. By uploading their resume against a target job description, students receive instant feedback on missing skills, enabling them to align their profiles before applying."
)

# --- CHAPTER 9 ---
doc.add_heading("CHAPTER 9 . Future Scope", level=1)

doc.add_paragraph(
    "Although the current implementation of ARSS is highly functional and responsive, several advanced features are planned for future versions:\n"
    "1. Large Language Model (LLM) Integration: Replacing or augmenting TF-IDF and spaCy with LLM APIs (e.g., Google Gemini or OpenAI GPT) will allow the system to perform deeper semantic understanding. LLMs can verify the quality of work experience and detect candidate achievements rather than relying on keyword similarity.\n"
    "2. Multilingual Support: Expanding the parsing capabilities by integrating multilingual spaCy models to screen resumes written in Spanish, French, German, or local languages.\n"
    "3. Optical Character Recognition (OCR): Utilizing libraries like Tesseract OCR to extract text from scanned, image-based PDF resumes, which currently return empty text strings.\n"
    "4. Automated Interview Scheduling: Integrating APIs for calendar services to automatically email qualified candidates and schedule preliminary technical screenings."
)

# --- CHAPTER 10 ---
doc.add_heading("CHAPTER 10. Conclusion", level=1)

doc.add_paragraph(
    "This project successfully designs, implements, and evaluates the AI Resume Screening System (ARSS). By integrating a modern web application "
    "with a structured Python NLP pipeline, the system bridges the gap between candidate accessibility and recruiter efficiency. The 5-step pipeline "
    "effectively automates parsing, information extraction, semantic document matching, hybrid scoring, and hard-gate classification. Crucially, the "
    "implementation of a persistent HTTP server for the Python pipeline solves the latency issues associated with loading large ML models, "
    "reducing processing time to under 500ms. Ultimately, the system establishes a scalable, objective, and high-performance foundation for "
    "automated talent acquisition, reducing hiring bias and operational overhead."
)

# --- REFERENCES ---
doc.add_heading("Reference", level=1)

p_ref1 = doc.add_paragraph()
p_ref1.paragraph_format.left_indent = Inches(0.5)
p_ref1.paragraph_format.first_line_indent = Inches(-0.5)
p_ref1.paragraph_format.space_after = Pt(8)
p_ref1.add_run("[1] Honnibal, M., & Montani, I. (2017). spaCy 2: Natural language understanding with Bloom embeddings, convolutional neural networks and incremental parsing. ")

p_ref2 = doc.add_paragraph()
p_ref2.paragraph_format.left_indent = Inches(0.5)
p_ref2.paragraph_format.first_line_indent = Inches(-0.5)
p_ref2.paragraph_format.space_after = Pt(8)
p_ref2.add_run("[2] Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. ")
p_ref2.add_run("Journal of Machine Learning Research, 12, 2825-2830.")

p_ref3 = doc.add_paragraph()
p_ref3.paragraph_format.left_indent = Inches(0.5)
p_ref3.paragraph_format.first_line_indent = Inches(-0.5)
p_ref3.paragraph_format.space_after = Pt(8)
p_ref3.add_run("[3] Salton, G., & Buckley, C. (1988). Term-weighting approaches in automatic text retrieval. ")
p_ref3.add_run("Information Processing & Management, 24(5), 513-523.")

p_ref4 = doc.add_paragraph()
p_ref4.paragraph_format.left_indent = Inches(0.5)
p_ref4.paragraph_format.first_line_indent = Inches(-0.5)
p_ref4.paragraph_format.space_after = Pt(8)
p_ref4.add_run("[4] Jurafsky, D., & Martin, J. H. (2020). Speech and Language Processing (3rd ed. draft). Stanford University.")

# Save the generated document
output_filename = "AI_Resume_Screening_System_Report.docx"
doc.save(output_filename)
print(f"Report successfully generated and saved to {output_filename}")

# Cleanup logo image
if os.path.exists(logo_path):
    os.remove(logo_path)
